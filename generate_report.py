import pandas as pd
from docx import Document
import tempfile
from adjustments import calculate_adjustments

def generate_report(df, subject_info, zillow_val=None, redfin_val=None, pdf_text="", real_avm=None):
    comps = []
    for _, row in df.iterrows():
        try:
            adj, ag_adj, ag_diff = calculate_adjustments(subject_info, row)
            close_price = row.get("Close Price", 0)
            concessions = row.get("Concessions", 0)
            ag_sf = row.get("AG SF", row.get("Above Grade Finished Area", 0))
            adjusted_price = close_price + concessions + adj
            ppsf = adjusted_price / ag_sf if ag_sf > 0 else 0
            comps.append({
                "Address": row.get("Street Address", row.get("Address", "N/A")),
                "Close Price": close_price,
                "Concessions": concessions,
                "AG SF": ag_sf,
                "AG Adj ($/sf)": 40,
                "AG Diff": ag_diff,
                "Total Adj": adj,
                "Adjusted Price": round(adjusted_price),
                "Adjusted PPSF": round(ppsf, 2)
            })
        except Exception as e:
            print(f"Error processing row: {e}")
            continue

    doc = Document()
    doc.add_heading("Market Valuation Report – Adjusted Comparison with Breakdown", 0)
    doc.add_paragraph("Date: July 17, 2025")
    doc.add_paragraph()
    doc.add_heading("Subject Property", level=1)
    doc.add_paragraph(f"Address: {subject_info.get('address', 'N/A')}")
    doc.add_paragraph(f"Above Grade SF: {subject_info.get('sqft', 0):,}")
    doc.add_paragraph(f"Bedrooms: {subject_info.get('bedrooms', 0)}")
    doc.add_paragraph(f"Bathrooms: {subject_info.get('bathrooms', 0)}")
    doc.add_paragraph()
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph("Close Price remains the original sale price. Adjustments are applied based on "
                      "square footage differences for Above Grade (AG) at $40/sf.")
    doc.add_paragraph()
    if comps:
        doc.add_heading("Comparable Properties Analysis", level=1)
        table = doc.add_table(rows=1, cols=len(comps[0]))
        table.style = 'Light Grid'
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(comps[0].keys()):
            hdr_cells[i].text = key
        for comp in comps:
            row_cells = table.add_row().cells
            for i, key in enumerate(comp.keys()):
                value = comp[key]
                if isinstance(value, (int, float)) and key != "Adjusted PPSF":
                    row_cells[i].text = f"{value:,}"
                else:
                    row_cells[i].text = str(value)

    doc.add_paragraph()
    doc.add_heading("Valuation Summary", level=1)
    if comps:
        avg_adjusted = sum(c['Adjusted Price'] for c in comps) / len(comps)
        avg_ppsf = sum(c['Adjusted PPSF'] for c in comps) / len(comps)
        doc.add_paragraph(f"Average Adjusted Price: ${avg_adjusted:,.0f}")
        doc.add_paragraph(f"Average Price Per SF: ${avg_ppsf:.2f}")
        estimates = []
        if zillow_val:
            estimates.append(zillow_val)
            doc.add_paragraph(f"Zillow Estimate: ${zillow_val:,}")
        if redfin_val:
            estimates.append(redfin_val)
            doc.add_paragraph(f"Redfin Estimate: ${redfin_val:,}")
        if real_avm:
            estimates.append(real_avm)
            doc.add_paragraph(f"RealAVM Estimate: ${real_avm:,}")
        if estimates:
            start_range = sum(estimates) / len(estimates)
            doc.add_paragraph()
            doc.add_paragraph(f"Market Range: ${start_range:,.0f} – ${int(avg_adjusted):,}")
            doc.add_paragraph("We've provided a valuation report using a consistent, data-driven adjustment model. "
                              "The range above reflects both public-facing AVMs and adjusted comparable sales analysis.")
    if pdf_text.strip():
        doc.add_paragraph()
        doc.add_heading("Extracted RealAVM Content", level=1)
        doc.add_paragraph(pdf_text[:1000] + "..." if len(pdf_text) > 1000 else pdf_text)
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    return temp.name
