import pandas as pd
import fitz
from docx import Document
from adjustments import calculate_adjustments
from datetime import date
import tempfile
import re

def parse_uploaded_csv(upload):
    df = pd.read_csv(upload) if upload.name.endswith("csv") else pd.read_excel(upload)
    rename_map = {
        "close price": "Close Price",
        "concessions": "Concessions",
        "above grade finished area": "AG SF",
        "basement total area": "Basement SF",
        "address": "Address"
    }
    df.columns = [c.strip().lower() for c in df.columns]
    df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
    df = df[["Address", "Close Price", "Concessions", "AG SF", "Basement SF"]]
    df = df.dropna(subset=["Close Price", "AG SF"])
    df["Concessions"] = df["Concessions"].fillna(0)
    return df

def extract_real_avm(upload, return_number=False):
    doc = fitz.open(stream=upload.read(), filetype="pdf")
    for page in doc:
        text = page.get_text()
        match = re.search(r"RealAVM™\s*\$([0-9,]+)", text)
        if match:
            return int(match.group(1).replace(",", ""))
    return None

def generate_report(df, subject_info, zillow, redfin, real_avm):
    ag_rate, basement_rate = 40, 10
    adjustments = []
    for _, row in df.iterrows():
        adj = calculate_adjustments(subject_info, row, ag_rate, basement_rate)
        adjustments.append({
            "Address": row["Address"],
            "Close Price": row["Close Price"],
            "Concessions": row["Concessions"],
            "AG SF": row["AG SF"],
            "Basement SF": row["Basement SF"],
            "AG Adj ($/sf)": ag_rate,
            "Basement Adj ($/sf)": basement_rate,
            "AG Diff": adj["AG Diff"],
            "Basement Diff": adj["Basement Diff"],
            "Adjusted Price": adj["Adjusted Price"],
            "Adjusted PPSF": adj["Adjusted PPSF"],
        })

    avg_adjusted = sum(x["Adjusted Price"] for x in adjustments) / len(adjustments)
    start_range = int((real_avm or 0 + zillow + redfin) / 3)

    doc = Document()
    doc.add_heading("Market Valuation Report – Adjusted Comparison with Breakdown", 0)
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(
        "Below is a breakdown of how adjustments were applied to comparable properties. "
        "Close Price remains the original sale price. Adjustments are applied based on square footage differences for Above Grade (AG) and Basement Finished area."
    )
    table = doc.add_table(rows=1, cols=len(adjustments[0]))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(adjustments[0].keys()):
        hdr_cells[i].text = col
    for adj in adjustments:
        row_cells = table.add_row().cells
        for i, val in enumerate(adj.values()):
            row_cells[i].text = f"${val:,.2f}" if isinstance(val, (int, float)) and i not in [3, 4, 5, 6, 7] else str(val)

    doc.add_paragraph()
    doc.add_heading("Finalized Listing Agent Blurb (with Value Range)", level=1)
    doc.add_paragraph(
        "We've provided a valuation report using a consistent, data-driven adjustment model designed for real-time market pricing—not retrospective appraisals. "
        "All comparable properties were selected within a 20% range of the subject's above-grade square footage and adjusted using a tiered system that distinguishes above-grade, basement, and finished basement areas separately. "
        "We also accounted for features like walkout basements, concessions, garage bays, and views."
    )
    doc.add_paragraph(
        f"Market Range: ${start_range:,.0f} – ${int(avg_adjusted):,}
"
        f"The starting value (${start_range:,}) reflects an average of public-facing automated estimates from RealAVM, Zillow, and Redfin.
"
        f"The ending value (${int(avg_adjusted):,}) is derived from the average of adjusted comparable sales after applying objective market adjustments."
    )

    output_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_path)
    return output_path