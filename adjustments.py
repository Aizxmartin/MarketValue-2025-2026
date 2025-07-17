def process_adjustments(df, subject_info, zillow, redfin, real_avm, pdf_text):
    from docx import Document
    from datetime import date
    import tempfile

    doc = Document()
    doc.add_heading("Market Valuation Report", 0)
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
    doc.add_heading("Subject Property", level=1)
    doc.add_paragraph(f"Address: {subject_info['address']}")
    doc.add_paragraph(f"Above Grade: {subject_info['sqft']} SqFt, Basement: {subject_info['basement']} SqFt, Total: {subject_info['sqft'] + subject_info['basement']} SqFt")
    doc.add_paragraph(f"Beds: {subject_info['beds']}, Baths: {subject_info['baths']}")
    doc.add_paragraph(f"Close Price: ${subject_info['price']:,.0f}")

    doc.add_heading("Online Valuations", level=1)
    avg_val = sum([v for v in [zillow, redfin, real_avm] if v]) / max(1, len([v for v in [zillow, redfin, real_avm] if v]))
    doc.add_paragraph(f"Zillow: ${zillow:,.0f}")
    doc.add_paragraph(f"Redfin: ${redfin:,.0f}")
    doc.add_paragraph(f"Real AVM: ${real_avm:,.0f}" if real_avm else "Real AVM: Not Extracted")
    doc.add_paragraph(f"Average of Online Estimates: ${avg_val:,.0f}")

    if pdf_text:
        doc.add_heading("Extracted Public Records Text", level=1)
        doc.add_paragraph(pdf_text)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name
